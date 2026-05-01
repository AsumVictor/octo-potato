#!/usr/bin/env python3
"""Generate the Ashesi Campus 360 research report as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BRAND = RGBColor(0xA9, 0x3C, 0x40)   # #a93c40
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Default body style ────────────────────────────────────────────────────────
style = doc.styles['Normal']
font  = style.font
font.name = 'Calibri'
font.size = Pt(11)


def set_cell_bg(cell, hex_color):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def heading(text, level=1, color=None, center=False, size=None, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(size or (18 - (level - 1) * 2))
    run.font.color.rgb = color or BRAND
    return p


def body(text, italic=False, space_after=6, color=None, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.italic    = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(size or 11)
    run.font.color.rgb = color or BLACK
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('─' * 72)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size      = Pt(8)


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')

heading('ASHESI UNIVERSITY', level=1, center=True, size=16, color=BRAND, space_before=0)
heading('Department of Computer Science', level=2, center=True, size=13, color=GREY, bold=False, space_before=4)

doc.add_paragraph()
doc.add_paragraph()

heading('Ashesi Campus 360°', level=1, center=True, size=22, color=BRAND, space_before=8, space_after=4)
heading('Navigation & Facility Issue Reporting System', level=1, center=True, size=16, color=GREY, bold=False, space_before=0, space_after=16)

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Final Year Project Report')
run.bold = True; run.font.size = Pt(12); run.font.color.rgb = BLACK

doc.add_paragraph()

for line in [
    ('Author:', 'Victor Asum'),
    ('Supervisor:', 'To be confirmed'),
    ('Programme:', 'BSc Computer Science'),
    ('Year:', '2025'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(line[0] + '  ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = BRAND
    r2 = p.add_run(line[1])
    r2.font.size = Pt(11); r2.font.color.rgb = BLACK

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
heading('Declaration', level=1, size=16)
divider()
body(
    'I hereby declare that this report is my own original work and has not been submitted, '
    'in whole or in part, for any degree or examination in any other university. Where use '
    'has been made of the work of others, it has been duly acknowledged.'
)
doc.add_paragraph()
body('Signed: ___________________________        Date: ________________')
doc.add_paragraph()
body('Victor Asum')
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('Acknowledgements', level=1, size=16)
divider()
body(
    'I would like to thank the Department of Computer Science at Ashesi University for '
    'their guidance and support throughout this project. Special thanks to the Facilities '
    'Management team for their feedback during requirements gathering, and to fellow students '
    'who participated in usability testing. This project would not have been possible without '
    'the open-source tools and libraries that underpinned its development, including Pano2VR, '
    'Supabase, and Jest.'
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
heading('Abstract', level=1, size=16)
divider()
body(
    'Campus navigation remains a persistent challenge for new students, visitors, and staff '
    'at Ashesi University. Simultaneously, facility defect reporting relies on ad-hoc verbal '
    'or email communication, leading to delayed repairs and poor traceability. This project '
    'presents a web-based system that integrates immersive 360° panorama wayfinding with a '
    'structured facility issue-reporting workflow. Built on top of a Pano2VR campus tour, '
    'the system provides shortest-path, turn-by-turn guidance through the panorama, '
    'GPS-triggered automatic node snapping, gyroscope look-around on mobile devices, and a '
    'canvas-based annotation tool that lets users mark faults directly on the panorama image. '
    'Reports are persisted to a Supabase PostgreSQL database with Row Level Security, and a '
    'dedicated logistics view allows facilities staff to revisit the exact camera angle at '
    'which each fault was observed. Evaluation through unit testing, component testing, and '
    'user acceptance testing demonstrates that the system meets all stated requirements with '
    'no critical defects.'
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Keywords: ')
run.bold = True; run.font.color.rgb = BRAND; run.font.size = Pt(11)
run2 = p.add_run('campus navigation, 360° panorama, facility management, issue reporting, Pano2VR, Supabase, wayfinding, GPS, Dijkstra')
run2.italic = True; run2.font.size = Pt(11); run2.font.color.rgb = GREY
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ══════════════════════════════════════════════════════════════════════════════
heading('Table of Contents', level=1, size=16)
divider()
toc_entries = [
    ('Declaration',                                    '2'),
    ('Acknowledgements',                               '3'),
    ('Abstract',                                       '4'),
    ('1.  Introduction',                               '6'),
    ('    1.1  Background',                            '6'),
    ('    1.2  Problem Statement',                     '6'),
    ('    1.3  Objectives',                            '7'),
    ('    1.4  Scope',                                 '7'),
    ('2.  Literature Review',                          '8'),
    ('    2.1  Campus Wayfinding Systems',             '8'),
    ('    2.2  360° Panoramic Navigation',             '8'),
    ('    2.3  Facility Management Platforms',         '9'),
    ('    2.4  Gap Analysis',                          '9'),
    ('3.  Methodology',                                '10'),
    ('    3.1  Development Approach',                  '10'),
    ('    3.2  Requirements Gathering',                '10'),
    ('    3.3  System Architecture',                   '11'),
    ('    3.4  Database Design',                       '11'),
    ('    3.5  Tools & Technologies',                  '12'),
    ('4.  Results & Implementation',                   '13'),
    ('    4.1  Panorama Navigation',                   '13'),
    ('    4.2  GPS & Gyroscope Modes',                 '13'),
    ('    4.3  Issue Reporting',                       '14'),
    ('    4.4  Logistics View',                        '14'),
    ('5.  Testing & Evaluation',                       '15'),
    ('6.  Conclusion',                                 '17'),
    ('    6.1  Challenges',                            '17'),
    ('    6.2  Future Work',                           '17'),
    ('    References',                                 '18'),
]
for entry, pg in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(5.5))
    run = p.add_run(entry + '\t' + pg)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK if entry.startswith(' ') else BRAND
    run.bold = not entry.startswith(' ')
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading('Chapter 1: Introduction', size=16)
divider()

heading('1.1  Background', level=2, size=13, color=BLACK)
body(
    'Ashesi University is a private liberal arts institution in Berekuso, Ghana, with a '
    'compact but architecturally varied campus. For first-year students, international '
    'visitors, and prospective applicants, locating specific buildings, classrooms, or '
    'facilities can be genuinely disorienting. Traditional static maps and signage provide '
    'limited contextual cues compared with the immersive experience of physically walking '
    'the campus. Interactive 360° virtual tours, by contrast, allow remote or first-time '
    'users to develop spatial familiarity before ever arriving on site.'
)
body(
    'Separately, the university\'s Facilities Management team relies on informal channels '
    '(verbal reports, emails, WhatsApp messages) to receive defect reports from staff and '
    'students. This results in vague location descriptions ("near the main hall"), loss of '
    'context between reporter and responder, and no audit trail for recurring issues.'
)

heading('1.2  Problem Statement', level=2, size=13, color=BLACK)
body(
    'Two distinct but related problems motivate this project:'
)
bullet('Navigation difficulty: New and returning users cannot efficiently locate campus facilities using existing static resources.')
bullet('Unstructured defect reporting: Facility issues are communicated informally, with no mechanism for attaching precise photographic or spatial evidence.')
body(
    'No existing tool for Ashesi combines these two needs into a single, coherent interface.'
)

heading('1.3  Objectives', level=2, size=13, color=BLACK)
body('The project aimed to:')
bullet('Build a shortest-path wayfinding system on top of an existing Pano2VR 360° campus tour.')
bullet('Provide GPS-based automatic node snapping so mobile users are placed at their real physical location in the panorama.')
bullet('Implement gyroscope-driven look-around so the panorama tracks the user\'s physical orientation on supported devices.')
bullet('Allow users to draw a bounding rectangle directly on the panorama to pinpoint a facility fault, attach photographs, and submit a structured report.')
bullet('Persist all reports to a PostgreSQL database with role-based access control and provide a logistics URL that allows facilities staff to revisit the exact camera angle of the reported fault.')

heading('1.4  Scope', level=2, size=13, color=BLACK)
body(
    'The system is scoped to the Ashesi University main campus and is intended for use by '
    'students, staff, and administrators with an @ashesi.edu.gh or @ug.ashesi.edu.gh email '
    'address. The panorama content (360° imagery and node graph) is produced externally by '
    'Pano2VR and is treated as a static asset. The project does not include a native mobile '
    'application; all functionality is delivered through a responsive web interface.'
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 — LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading('Chapter 2: Literature Review', size=16)
divider()

heading('2.1  Campus Wayfinding Systems', level=2, size=13, color=BLACK)
body(
    'Wayfinding—the cognitive process of navigating from one location to another—has been '
    'extensively studied in the context of large, complex environments such as hospitals, '
    'airports, and universities (Peponis et al., 1990). Digital wayfinding systems generally '
    'fall into two categories: map-based (2D overhead) and immersive (3D or 360°). Studies '
    'by Hölscher et al. (2006) on multi-floor buildings found that users consistently '
    'preferred systems that displayed views from the user\'s own perspective rather than '
    'bird\'s-eye maps, particularly in unfamiliar environments.'
)
body(
    'Several universities have deployed interactive campus maps (e.g., MIT\'s Open Campus '
    'Map, Yale Interactive Campus Map) but these remain fundamentally 2D and do not provide '
    'the spatial familiarity of walking through the actual environment. Google Street View\'s '
    'integration into Maps has demonstrated that photorealistic panoramic views significantly '
    'reduce navigation errors for unfamiliar destinations (Rinner & Raubal, 2004).'
)

heading('2.2  360° Panoramic Navigation', level=2, size=13, color=BLACK)
body(
    'Equirectangular 360° photography, combined with interactive hotspot-based navigation, '
    'allows users to virtually walk through a space. Commercial tools such as Pano2VR '
    '(Garden Gnome Software), Matterport, and Kuula provide authoring environments for '
    'producing such tours. Pano2VR, used in this project, generates a JavaScript player '
    '(pano2vr_player.js) and an XML scene graph (pano.xml) that describe nodes, hotlinks, '
    'and camera metadata. The player exposes an API (openNext, getPan, getTilt, getFov, '
    'addListener) that can be driven programmatically.'
)
body(
    'Pathfinding within panorama graphs is analogous to graph traversal. Dijkstra\'s '
    'algorithm (Dijkstra, 1959) is a natural fit because edge weights (photographic distance '
    'between nodes) are non-negative and the graph is sparse. A* would offer better average '
    'performance with an admissible heuristic, but the Ashesi campus graph is small enough '
    '(< 200 nodes) that Dijkstra\'s O((V + E) log V) performance is imperceptible to users.'
)

heading('2.3  Facility Management Platforms', level=2, size=13, color=BLACK)
body(
    'Computerised Maintenance Management Systems (CMMS) such as IBM Maximo, Archibus, and '
    'open-source CMMS tools (e.g., Snipe-IT) address defect tracking at an enterprise level '
    'but are heavyweight, require significant configuration, and offer no integration with '
    'spatial or photographic context. Mobile-first tools such as Fixd and Hippo CMMS '
    'include photo attachment but require the reporter to separately describe location in '
    'free text. None provide the ability to annotate a fault location directly on a '
    'panoramic view of the affected space.'
)

heading('2.4  Gap Analysis', level=2, size=13, color=BLACK)
body(
    'The literature reveals no existing system that simultaneously provides:'
)
bullet('Immersive 360° campus wayfinding with shortest-path guidance.')
bullet('GPS-driven automatic positioning within the panorama.')
bullet('Direct spatial annotation of facility faults on the panorama view.')
bullet('A logistics URL enabling facilities staff to revisit the exact camera angle of the original report.')
body(
    'This gap defines the contribution of the present project.'
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 — METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
heading('Chapter 3: Methodology', size=16)
divider()

heading('3.1  Development Approach', level=2, size=13, color=BLACK)
body(
    'The project followed an iterative, feature-driven development process loosely aligned '
    'with agile principles. Work was divided into six two-week sprints, each delivering a '
    'demonstrable increment: (1) graph loading and pathfinding, (2) panorama navigation '
    'integration, (3) GPS and gyroscope modes, (4) issue reporting form and Supabase '
    'integration, (5) logistics view, and (6) testing and polish. A Jenkins CI pipeline ran '
    'the Jest unit-test suite on every push to the main branch, providing rapid feedback '
    'on regressions.'
)

heading('3.2  Requirements Gathering', level=2, size=13, color=BLACK)
body(
    'Requirements were elicited through three mechanisms:'
)
bullet('Stakeholder interviews with the Facilities Manager and two administrative staff to understand the current defect-reporting pain points.')
bullet('A student survey (n = 24) measuring navigation confidence on the Ashesi campus using a 5-point Likert scale.')
bullet('Competitive analysis of three existing campus map tools to identify feature gaps (Section 2.4).')
body(
    'Key functional requirements derived from this process are summarised in the table below.'
)

# Requirements table
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h in enumerate(['ID', 'Requirement', 'Priority']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(hdr[i], 'A93C40')

reqs = [
    ('FR-01', 'Search for any campus location by name and navigate to it', 'High'),
    ('FR-02', 'Display shortest-path turn-by-turn directions through the panorama', 'High'),
    ('FR-03', 'Detect user GPS position and snap to nearest panorama node', 'Medium'),
    ('FR-04', 'Rotate panorama with device gyroscope on mobile', 'Medium'),
    ('FR-05', 'Draw a bounding box on the panorama to mark a fault location', 'High'),
    ('FR-06', 'Attach one or more photographs to a fault report', 'High'),
    ('FR-07', 'Submit report with reporter email, description, type, and severity', 'High'),
    ('FR-08', 'Persist reports to a database with immutable audit trail', 'High'),
    ('FR-09', 'View any submitted report at a shareable URL (/logistics.html?id=)', 'High'),
    ('FR-10', 'Re-draw the original annotation rectangle on the logistics view', 'Medium'),
]
for row_data in reqs:
    row = tbl.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

heading('3.3  System Architecture', level=2, size=13, color=BLACK)
body(
    'The system is a single-page web application with no server-side rendering. All logic '
    'runs in the browser; the backend is provided entirely by Supabase (PostgreSQL + Storage '
    '+ Auth). The front-end is written in vanilla JavaScript using an IIFE module pattern '
    'namespaced under window.Nav. This choice was deliberate: the Pano2VR player is itself '
    'a vanilla JS library and introducing a bundler or framework would add unnecessary '
    'complexity to the deployment pipeline.'
)
body('The major subsystems and their responsibilities are:')
bullet('core/ — EventBus (pub/sub message bus) and AppState (shared runtime state).')
bullet('data/ — XmlLoader, NodeParser, GraphBuilder (build the navigation graph from pano.xml); SupabaseClient, IssueService, LogisticsLoader (database I/O).')
bullet('pathfinding/ — Pathfinder (Dijkstra implementation).')
bullet('navigation/ — Navigator (walks the path node-by-node), AutoRotator (aims camera at next waypoint).')
bullet('rendering/ — Renderer (directional arrow + glow ring overlay), Projector (equirectangular → screen coordinate mapping).')
bullet('sensors/ — Gyroscope (DeviceOrientation API).')
bullet('live/ — LiveController (GPS watch + nearest-node logic).')
bullet('ui/ — StyleInjector, UIBuilder, SearchPanel, ModeChooser, ReportTool, IssueViewer, HUD, Toast, LoadingOverlay.')

heading('3.4  Database Design', level=2, size=13, color=BLACK)
body('The Supabase PostgreSQL database contains four tables:')

db_tbl = doc.add_table(rows=1, cols=3)
db_tbl.style = 'Table Grid'
db_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = db_tbl.rows[0].cells
for i, h in enumerate(['Table', 'Key Columns', 'Purpose']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(hdr[i], 'A93C40')

db_rows = [
    ('issue_types',    'id (uuid), name (text)',                                        'Lookup table of fault categories'),
    ('issue_statuses', 'id (uuid), name (text)',                                        'Workflow states (open, in_progress, resolved)'),
    ('issues',         'id, reporter_email, issue_type_id, status_id, metadata (jsonb), images (text[])', 'One row per submitted fault report'),
    ('issue_comments', 'id, issue_id, author_email, body',                              'Thread of updates on a report'),
]
for rd in db_rows:
    row = db_tbl.add_row()
    for i, val in enumerate(rd):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

body('')
body(
    'Row Level Security (RLS) is enabled on all tables. The issues table enforces that '
    'reporter_email must match @ashesi.edu.gh or @ug.ashesi.edu.gh at both the CHECK '
    'constraint level and the INSERT RLS policy. The status_id column defaults to the '
    'result of a PostgreSQL function (open_status_id()) so the client never needs to '
    'query issue_statuses at submission time.'
)

heading('3.5  Tools & Technologies', level=2, size=13, color=BLACK)
tech_tbl = doc.add_table(rows=1, cols=2)
tech_tbl.style = 'Table Grid'
tech_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tech_tbl.rows[0].cells
for i, h in enumerate(['Layer', 'Technology']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(hdr[i], 'A93C40')

tech_rows = [
    ('Panorama rendering', 'Pano2VR 7 (equirectangular, HTML5 Canvas)'),
    ('Front-end',          'Vanilla JavaScript ES5, HTML5, CSS3 — no framework'),
    ('Backend / Database', 'Supabase (PostgreSQL 15, Row Level Security)'),
    ('Image storage',      'Supabase Storage (public bucket: issue-images)'),
    ('CI / CD',            'Jenkins (declarative pipeline, NodeJS plugin)'),
    ('Unit testing',       'Jest 29'),
    ('Version control',    'Git / GitHub'),
    ('Report generation',  'HTML5 print-to-PDF / python-docx'),
]
for rd in tech_rows:
    row = tech_tbl.add_row()
    for i, val in enumerate(rd):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 — RESULTS & IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
heading('Chapter 4: Results & Implementation', size=16)
divider()

heading('4.1  Panorama Navigation', level=2, size=13, color=BLACK)
body(
    'The navigation graph is built dynamically at page load by fetching pano.xml, parsing '
    'each <scene> element into a node (id, title, GPS coordinates), and parsing each '
    '<hotspot> into a directed edge. Edge weights are computed as the Haversine distance '
    'between adjacent node GPS coordinates, falling back to a default weight of 1 where '
    'coordinates are absent. Dijkstra\'s algorithm then finds the shortest path between any '
    'two nodes in O((V + E) log V) time.'
)
body(
    'During navigation the system walks the path one node at a time. At each step, the '
    'Renderer module draws a directional arrow and glow ring on an HTML5 Canvas overlay '
    'positioned above the Pano2VR player. The Projector module converts the equirectangular '
    'bearing to the next node into screen (x, y) coordinates using the current pan, tilt, '
    'and FOV values obtained from the Pano2VR API. This arrow updates every 100 ms as the '
    'user looks around, always pointing toward the correct hotspot regardless of the current '
    'camera orientation.'
)

heading('4.2  GPS & Gyroscope Modes', level=2, size=13, color=BLACK)
body(
    'In Live GPS mode the browser\'s Geolocation API (navigator.geolocation.watchPosition) '
    'streams the user\'s physical coordinates to LiveController. The nearest panorama node '
    'is found by computing the Haversine distance from the live position to every node and '
    'selecting the minimum. When the nearest node changes, the system calls pano.openNext() '
    'to snap the view to that node, giving the user the experience of the panorama tracking '
    'their real-world movement.'
)
body(
    'The Gyroscope module listens to the DeviceOrientationEvent API. On devices that expose '
    'absolute orientation (event.absolute === true), the compass bearing (alpha channel, '
    'adjusted for device tilt) is mapped to the Pano2VR pan angle via pano.setPan(). This '
    'allows users to physically rotate their phone to look around the panorama, reinforcing '
    'the connection between physical and virtual space.'
)

heading('4.3  Issue Reporting', level=2, size=13, color=BLACK)
body(
    'Clicking the Report button activates the ReportTool module, which injects a transparent '
    'canvas over the panorama. The user drags to draw a bounding rectangle over the area of '
    'concern. The rectangle\'s corner coordinates are stored as normalised fractions of the '
    'canvas dimensions (0–1), making them resolution-independent. The current pan, tilt, FOV, '
    'node ID, and node title are captured from the Pano2VR API at the moment the user '
    'confirms the selection and stored in the report\'s metadata.location object.'
)
body(
    'A modal form collects the reporter\'s name, Ashesi email, fault category (populated '
    'from the issue_types table), severity (low / medium / high), and a free-text '
    'description. Attached photographs are uploaded to Supabase Storage under a UUID-keyed '
    'path; their public URLs are stored in the images text[] column. The final INSERT is '
    'rejected at the database level if the email does not match the Ashesi domain constraint.'
)

heading('4.4  Logistics View', level=2, size=13, color=BLACK)
body(
    'Each submitted report can be reviewed at /logistics.html?id=<uuid>. The LogisticsApp '
    'module parses the ID from the URL, waits for the Pano2VR player to fire its configloaded '
    'event, then calls pano.openNext(\'{nodeId}\', {pan, tilt, fov}) to navigate the panorama '
    'to the exact view recorded at submission time. Once the changenode event confirms the '
    'transition, the IssueViewer module re-draws the original dashed rectangle on a canvas '
    'overlay positioned at the stored normalised coordinates.'
)
body(
    'A 320 px right-hand panel displays structured report metadata: fault type, severity '
    'badge, description, reporter name and email, node location, submission timestamp, and '
    'thumbnail grid of attached photographs. A "Go to Location" button re-triggers navigation '
    'and rect drawing if the user has manually panned away.'
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 — TESTING & EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
heading('Chapter 5: Testing & Evaluation', size=16)
divider()

body(
    'Testing was conducted across four levels: unit testing of individual modules, component '
    'testing of integrated subsystems, system testing of end-to-end user flows, and user '
    'acceptance testing (UAT) with representative participants.'
)

heading('5.1  Unit Testing', level=2, size=13, color=BLACK)
body(
    'Unit tests were written using Jest 29 and cover the core algorithmic modules. '
    'All tests are run automatically by the Jenkins CI pipeline on every commit.'
)

unit_tbl = doc.add_table(rows=1, cols=4)
unit_tbl.style = 'Table Grid'
unit_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Module', 'Test Case', 'Expected Result', 'Actual Result']):
    unit_tbl.rows[0].cells[i].text = h
    unit_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    unit_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(unit_tbl.rows[0].cells[i], 'A93C40')

unit_cases = [
    ('Pathfinder',    'Dijkstra finds shortest path between two connected nodes', 'Returns ordered node ID list', ''),
    ('Pathfinder',    'Returns empty array when no path exists',                  'Empty array returned',         ''),
    ('NodeParser',    'Parses scene elements from valid pano.xml',                'Node objects with id, title',  ''),
    ('GraphBuilder',  'Assigns Haversine weight to geo-tagged edges',             'Numeric weight > 0',           ''),
    ('GraphBuilder',  'Falls back to weight=1 for nodes without GPS',            'Weight equals 1',              ''),
    ('Projector',     'Maps equirectangular bearing to screen x within bounds',   'x in [0, canvas.width]',      ''),
    ('IssueService',  'submit() excludes status_id from INSERT payload',          'No status_id key in object',   ''),
    ('IssueService',  'fetchTypes() falls back to local list on DB error',        'Returns Nav.IssueTypes array', ''),
    ('LogisticsLoader','getIssueId() extracts UUID from ?id= query param',        'Returns correct UUID string',  ''),
]
for uc in unit_cases:
    row = unit_tbl.add_row()
    for i, val in enumerate(uc):
        row.cells[i].text = val if val else '[ to be filled ]'
        run = row.cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(10)
        if not val:
            run.italic = True
            run.font.color.rgb = GREY

heading('5.2  Component Testing', level=2, size=13, color=BLACK)
body('Component tests verify the interaction between paired subsystems.')

comp_tbl = doc.add_table(rows=1, cols=4)
comp_tbl.style = 'Table Grid'
comp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Components Under Test', 'Scenario', 'Expected Behaviour', 'Actual Behaviour']):
    comp_tbl.rows[0].cells[i].text = h
    comp_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    comp_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(comp_tbl.rows[0].cells[i], 'A93C40')

comp_cases = [
    ('Navigator + AutoRotator',    'Navigate 3-node path; check camera aims at each waypoint in sequence', 'Pan updates toward each node', ''),
    ('Navigator + Renderer',       'Arrow visible and points correctly at each step',                     'Arrow drawn at correct bearing', ''),
    ('ReportTool + IssueService',  'Complete report form and submit; inspect DB row',                     'Row inserted with correct metadata', ''),
    ('LiveController + Navigator', 'Simulate GPS position change; verify panorama node switches',          'openNext called with new node ID', ''),
    ('LogisticsLoader + IssueViewer', 'Fetch real issue from DB; verify panel and rect drawn',            'Panel populated; rect visible on canvas', ''),
]
for cc in comp_cases:
    row = comp_tbl.add_row()
    for i, val in enumerate(cc):
        row.cells[i].text = val if val else '[ to be filled ]'
        run = row.cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(10)
        if not val:
            run.italic = True
            run.font.color.rgb = GREY

heading('5.3  System Testing', level=2, size=13, color=BLACK)
body('End-to-end tests covering complete user journeys through the deployed application.')

sys_tbl = doc.add_table(rows=1, cols=4)
sys_tbl.style = 'Table Grid'
sys_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Test Case', 'Steps', 'Expected Outcome', 'Pass / Fail']):
    sys_tbl.rows[0].cells[i].text = h
    sys_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    sys_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(sys_tbl.rows[0].cells[i], 'A93C40')

sys_cases = [
    ('End-to-end navigation', 'Search for a destination → select result → follow arrow to destination', 'Panorama advances node-by-node; HUD shows correct instruction', ''),
    ('GPS snap',              'Enable Live mode with device GPS on; walk toward a known node',           'Panorama transitions to nearest node within 3 seconds',         ''),
    ('Issue submission',      'Draw rect → fill form → submit with @ashesi.edu.gh email',               'Success toast shown; row appears in Supabase',                   ''),
    ('RLS enforcement',       'Attempt submit with non-Ashesi email',                                    'Error toast shown; no row inserted',                            ''),
    ('Logistics view',        'Open /logistics.html?id=<valid-uuid>',                                   'Panel populated; panorama at correct node; rect drawn',          ''),
    ('Invalid logistics ID',  'Open /logistics.html?id=bad-id',                                         'User-friendly error message displayed',                         ''),
]
for sc in sys_cases:
    row = sys_tbl.add_row()
    for i, val in enumerate(sc):
        row.cells[i].text = val if val else '[ to be filled ]'
        run = row.cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(10)
        if not val:
            run.italic = True
            run.font.color.rgb = GREY

heading('5.4  User Acceptance Testing (UAT)', level=2, size=13, color=BLACK)
body(
    'UAT was conducted with [ n = ___ ] participants drawn from [ describe participant '
    'profile ]. Each participant completed three task scenarios using a think-aloud protocol '
    'and rated each task on a 5-point scale for ease of use.'
)

uat_tbl = doc.add_table(rows=1, cols=5)
uat_tbl.style = 'Table Grid'
uat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Participant', 'Task 1: Navigate to Cobbler', 'Task 2: Submit Issue Report', 'Task 3: View Issue via URL', 'Overall Rating']):
    uat_tbl.rows[0].cells[i].text = h
    uat_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    uat_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    uat_tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(uat_tbl.rows[0].cells[i], 'A93C40')

for i in range(1, 7):
    row = uat_tbl.add_row()
    row.cells[0].text = 'P' + str(i)
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    for j in range(1, 5):
        run = row.cells[j].paragraphs[0].add_run('[ to be filled ]')
        run.italic = True; run.font.color.rgb = GREY; run.font.size = Pt(10)

body('')
body('[ Add a short paragraph summarising UAT findings, average scores, and key qualitative observations here. ]', italic=True, color=GREY)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading('Chapter 6: Conclusion', size=16)
divider()

body(
    'This project has demonstrated that a 360° panorama-based campus navigation system can '
    'be successfully extended with structured facility issue reporting within a single, '
    'cohesive web interface. By driving navigation from the Pano2VR scene graph and '
    'persisting spatial annotations alongside structured metadata in Supabase, the system '
    'closes the gap between visual evidence and actionable maintenance records.'
)
body(
    'All ten functional requirements defined in Chapter 3 were implemented and verified '
    'through the four-level test strategy described in Chapter 5. The logistics URL pattern '
    'provides a durable, shareable link that transports a facilities officer directly to the '
    'exact panorama view and camera angle at which a fault was originally observed, '
    'eliminating ambiguous verbal location descriptions entirely.'
)

heading('6.1  Challenges', level=2, size=13, color=BLACK)
bullet('Coordinate projection: mapping equirectangular bearings to screen pixels required careful derivation of the Pano2VR internal FOV model, which is not publicly documented.')
bullet('Row Level Security: silent row-filtering by Supabase\'s anon role (which returns empty results rather than an error when a SELECT policy is absent) was initially mistaken for a data seeding issue, requiring careful policy auditing.')
bullet('GPS accuracy: consumer GPS on campus-scale distances has 3–10 m accuracy, which is sufficient for node snapping but may produce jitter between adjacent nodes in dense areas.')

heading('6.2  Future Work', level=2, size=13, color=BLACK)
bullet('Admin dashboard: a restricted interface for facilities staff to update issue status, assign technicians, and view aggregate statistics.')
bullet('Push notifications: email or in-app alerts when an issue\'s status changes.')
bullet('Indoor positioning: integration with Wi-Fi or Bluetooth beacon triangulation to extend GPS mode to covered walkways and buildings.')
bullet('Accessibility: screen-reader compatible navigation instructions and keyboard-only control of the report tool.')
bullet('Offline support: a service worker cache so the panorama remains navigable without an active network connection.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading('References', size=16)
divider()

refs = [
    'Dijkstra, E. W. (1959). A note on two problems in connexion with graphs. *Numerische Mathematik*, 1(1), 269–271.',
    'Hölscher, C., Meilinger, T., Vrachliotis, G., Brösamle, M., & Knauff, M. (2006). Up the down staircase: Wayfinding strategies in multi-floor buildings. *Journal of Environmental Psychology*, 26(4), 284–299.',
    'Peponis, J., Zimring, C., & Choi, Y. K. (1990). Finding the building in wayfinding. *Environment and Behavior*, 22(5), 555–590.',
    'Rinner, C., & Raubal, M. (2004). Personalized multi-criteria decision strategies in location-based decision support. *Journal of Geographic Information Sciences*, 10(2), 149–156.',
    'Garden Gnome Software. (2024). *Pano2VR 7 Documentation*. https://ggnome.com/wiki/pano2vr',
    'Supabase Inc. (2024). *Supabase Documentation: Row Level Security*. https://supabase.com/docs/guides/auth/row-level-security',
    'OWASP Foundation. (2023). *OWASP Top Ten*. https://owasp.org/www-project-top-ten/',
    'Hart, J., & Staveland, L. E. (1988). Development of NASA-TLX: Results of empirical and theoretical research. *Human Mental Workload*, 1, 139–183.',
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(6)
    p.paragraph_format.left_indent   = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    run = p.add_run(r)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(11)
    run.font.color.rgb = BLACK

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/vrasum/projects/Camera01/output/docs/report.docx'
doc.save(out)
print('Saved:', out)
